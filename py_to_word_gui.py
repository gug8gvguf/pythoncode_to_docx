import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob


class PyToWordGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Python代码转Word工具")
        self.root.geometry("700x500")

        self.py_files = []
        self.output_dir = ""

        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill=tk.X, pady=5)

        ttk.Label(top_frame, text="选择Python文件:").pack(side=tk.LEFT, padx=5)
        self.add_files_btn = ttk.Button(top_frame, text="添加文件", command=self.add_files)
        self.add_files_btn.pack(side=tk.LEFT, padx=5)
        self.add_dir_btn = ttk.Button(top_frame, text="添加目录", command=self.add_dir)
        self.add_dir_btn.pack(side=tk.LEFT, padx=5)
        self.clear_files_btn = ttk.Button(top_frame, text="清空列表", command=self.clear_files)
        self.clear_files_btn.pack(side=tk.LEFT, padx=5)

        list_frame = ttk.LabelFrame(main_frame, text="文件列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.files_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, selectmode=tk.EXTENDED)
        self.files_listbox.pack(fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.files_listbox.yview)

        self.remove_file_btn = ttk.Button(list_frame, text="移除选中文件", command=self.remove_selected)
        self.remove_file_btn.pack(fill=tk.X, pady=5)

        output_frame = ttk.LabelFrame(main_frame, text="输出设置", padding=10)
        output_frame.pack(fill=tk.X, pady=5)

        output_row = ttk.Frame(output_frame)
        output_row.pack(fill=tk.X, pady=5)

        ttk.Label(output_row, text="输出目录:").pack(side=tk.LEFT, padx=5)
        self.output_entry = ttk.Entry(output_row)
        self.output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.browse_output_btn = ttk.Button(output_row, text="浏览", command=self.browse_output)
        self.browse_output_btn.pack(side=tk.LEFT, padx=5)

        options_frame = ttk.LabelFrame(main_frame, text="选项", padding=10)
        options_frame.pack(fill=tk.X, pady=5)

        options_row = ttk.Frame(options_frame)
        options_row.pack(fill=tk.X, pady=5)

        self.show_line_numbers = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_row, text="显示行号", variable=self.show_line_numbers).pack(side=tk.LEFT, padx=10)

        self.syntax_highlight = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_row, text="语法高亮", variable=self.syntax_highlight).pack(side=tk.LEFT, padx=10)

        self.merge_to_one = tk.BooleanVar(value=False)
        ttk.Checkbutton(options_row, text="多合一(合并到一个文档)", variable=self.merge_to_one).pack(side=tk.LEFT, padx=10)

        self.show_file_path = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_row, text="显示文件路径", variable=self.show_file_path).pack(side=tk.LEFT, padx=10)

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        self.convert_btn = ttk.Button(btn_frame, text="开始转换", command=self.convert)
        self.convert_btn.pack(side=tk.LEFT, padx=5)

        self.status_var = tk.StringVar(value="就绪")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X, pady=5)

    def add_files(self):
        file_paths = filedialog.askopenfilenames(
            title="选择Python文件",
            filetypes=[("Python文件", "*.py"), ("所有文件", "*.*")]
        )
        if file_paths:
            for file_path in file_paths:
                if file_path not in self.py_files and file_path.endswith('.py'):
                    self.py_files.append(file_path)
                    self.files_listbox.insert(tk.END, os.path.basename(file_path))
            self.update_status()

    def add_dir(self):
        dir_path = filedialog.askdirectory(title="选择目录")
        if dir_path:
            py_files = glob.glob(os.path.join(dir_path, '*.py'))
            if not py_files:
                messagebox.showwarning("警告", "目录中没有Python文件")
                return

            count = 0
            for file_path in py_files:
                if file_path not in self.py_files:
                    self.py_files.append(file_path)
                    self.files_listbox.insert(tk.END, os.path.basename(file_path))
                    count += 1

            if count > 0:
                messagebox.showinfo("提示", f"成功添加 {count} 个文件")
                self.update_status()
            else:
                messagebox.showinfo("提示", "没有新文件可添加")

    def clear_files(self):
        self.py_files = []
        self.files_listbox.delete(0, tk.END)
        self.update_status()

    def remove_selected(self):
        selected_indices = self.files_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("警告", "请先选择要移除的文件")
            return

        for index in reversed(selected_indices):
            self.files_listbox.delete(index)
            del self.py_files[index]

        self.update_status()

    def browse_output(self):
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_dir = dir_path
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, dir_path)

    def update_status(self):
        self.status_var.set(f"已选择 {len(self.py_files)} 个文件")

    def py_to_word(self, py_file_path, output_docx_path=None):
        if not os.path.exists(py_file_path):
            return False, f"文件不存在: {py_file_path}"

        if not py_file_path.endswith('.py'):
            return False, f"不是Python文件: {py_file_path}"

        if output_docx_path is None:
            base_name = os.path.splitext(py_file_path)[0]
            output_docx_path = f"{base_name}.docx"

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)

        title = os.path.basename(py_file_path)
        doc.add_heading(title, level=1)

        with open(py_file_path, 'r', encoding='utf-8') as f:
            code_lines = f.readlines()

        code_paragraph = doc.add_paragraph()
        code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for line_num, line in enumerate(code_lines, 1):
            prefix = f"{line_num:4d}  " if self.show_line_numbers.get() else ""
            run = code_paragraph.add_run(f"{prefix}{line}")
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 51, 102)

            if self.syntax_highlight.get():
                if line.strip().startswith('#'):
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif 'import ' in line or 'from ' in line:
                    run.font.color.rgb = RGBColor(128, 0, 128)
                elif 'def ' in line or 'class ' in line:
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif 'if ' in line or 'elif ' in line or 'else:' in line:
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif 'for ' in line or 'while ' in line:
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif 'return ' in line:
                    run.font.color.rgb = RGBColor(0, 0, 255)

        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        try:
            doc.save(output_docx_path)
            return True, f"成功: {output_docx_path}"
        except Exception as e:
            return False, f"保存失败: {str(e)}"

    def convert(self):
        if not self.py_files:
            messagebox.showwarning("警告", "请先选择Python文件")
            return

        if self.merge_to_one.get():
            self.merge_convert()
        else:
            self.batch_convert()

    def batch_convert(self):
        if not self.output_dir:
            self.output_dir = os.path.dirname(self.py_files[0])
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, self.output_dir)

        success_count = 0
        fail_count = 0
        fail_messages = []

        self.status_var.set("正在转换...")
        self.root.update()

        for i, py_file in enumerate(self.py_files, 1):
            base_name = os.path.splitext(os.path.basename(py_file))[0]
            output_path = os.path.join(self.output_dir, f"{base_name}.docx")

            success, message = self.py_to_word(py_file, output_path)
            self.status_var.set(f"转换中 ({i}/{len(self.py_files)})...")
            self.root.update()

            if success:
                success_count += 1
            else:
                fail_count += 1
                fail_messages.append(f"  - {os.path.basename(py_file)}: {message}")

        self.status_var.set(f"转换完成! 成功: {success_count}, 失败: {fail_count}")

        if fail_count > 0:
            error_msg = "\n".join(fail_messages)
            messagebox.showwarning("部分失败", f"转换完成!\n成功: {success_count}\n失败: {fail_count}\n\n失败详情:\n{error_msg}")
        else:
            messagebox.showinfo("成功", f"全部转换成功!\n共 {success_count} 个文件\n\n输出目录:\n{self.output_dir}")

    def merge_convert(self):
        output_file = filedialog.asksaveasfilename(
            title="保存合并文档",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )

        if not output_file:
            return

        self.status_var.set("正在合并文档...")
        self.root.update()

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)

        doc.add_heading('Python代码合集', level=1)
        doc.add_paragraph(f'共 {len(self.py_files)} 个文件', style='Heading 2')

        success_count = 0
        fail_count = 0

        for idx, py_file in enumerate(self.py_files, 1):
            if not os.path.exists(py_file):
                fail_count += 1
                continue

            if not py_file.endswith('.py'):
                fail_count += 1
                continue

            doc.add_page_break()

            file_name = os.path.basename(py_file)
            doc.add_heading(f'{idx}. {file_name}', level=2)
            if self.show_file_path.get():
                doc.add_paragraph(f'文件路径: {py_file}')

            with open(py_file, 'r', encoding='utf-8') as f:
                code_lines = f.readlines()

            code_paragraph = doc.add_paragraph()
            code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for line_num, line in enumerate(code_lines, 1):
                prefix = f"{line_num:4d}  " if self.show_line_numbers.get() else ""
                run = code_paragraph.add_run(f"{prefix}{line}")
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 51, 102)

                if self.syntax_highlight.get():
                    if line.strip().startswith('#'):
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif 'import ' in line or 'from ' in line:
                        run.font.color.rgb = RGBColor(128, 0, 128)
                    elif 'def ' in line or 'class ' in line:
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif 'if ' in line or 'elif ' in line or 'else:' in line:
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif 'for ' in line or 'while ' in line:
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif 'return ' in line:
                        run.font.color.rgb = RGBColor(0, 0, 255)

            success_count += 1
            self.status_var.set(f"合并中 ({idx}/{len(self.py_files)})...")
            self.root.update()

        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        try:
            doc.save(output_file)
            self.status_var.set(f"合并完成! 成功: {success_count}, 失败: {fail_count}")
            messagebox.showinfo("成功", f"合并文档已保存!\n共 {success_count} 个文件\n\n文件路径:\n{output_file}")
        except Exception as e:
            self.status_var.set("合并失败")
            messagebox.showerror("错误", f"合并失败: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = PyToWordGUI(root)
    root.mainloop()
